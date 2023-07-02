from docx import Document
from docx.shared import RGBColor
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml

# Load the Word document
doc = Document('sample.docx')

# Define the color for highlighting the first row (e.g., light blue)
first_row_highlight_color = RGBColor(173, 216, 230)  # Light blue highlight for the first row

# Define the color for highlighting empty cells (e.g., yellow)
empty_cell_highlight_color = RGBColor(255, 255, 0)  # Yellow highlight for empty cells

# Iterate over the tables in the document
for table in doc.tables:
    # Get the first row of the table
    first_row = table.rows[0]

    # Check if any cell in the first row is not bold
    if any(not cell.paragraphs[0].runs[0].bold for cell in first_row.cells):
        # Highlight the first row with the specified color
        for cell in first_row.cells:
            shading = parse_xml('<w:shd {} w:fill="ADD8E6"/>'.format(nsdecls('w')))
            cell._element.get_or_add_tcPr().append(shading)

    # Process each row in the table
    for row in table.rows:
        # Process each cell in the row
        for cell in row.cells:
            # Check if the cell is empty
            if cell.text.strip() == "":
                # Set the background color of the empty cell to yellow
                shading = parse_xml('<w:shd {} w:fill="FFFF00"/>'.format(nsdecls('w')))
                cell._tc.get_or_add_tcPr().append(shading)

# Save the modified document with highlighted first rows and empty cells
doc.save('modified_document.docx')
