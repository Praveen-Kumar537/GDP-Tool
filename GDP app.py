import tkinter as tk
from tkinter import filedialog
import os
import docx
from docx.shared import RGBColor
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml
from docx.enum.text import WD_COLOR_INDEX

def upload_file():
    file_path = filedialog.askopenfilename(filetypes=[("DOCX files", "*.docx")])
    if file_path:
        file_entry.delete(0, tk.END)
        file_entry.insert(0, file_path)

def modify_docx():
    file_path = file_entry.get()
    if not file_path:
        status_label.config(text="Please upload a DOCX file.", fg="red")
        return

    doc = docx.Document(file_path)

    # Add your modifications here
    # For example, you can add a new paragraph:
    new_paragraph = "This is a new paragraph added by the program."
    doc.add_paragraph(new_paragraph)

    for table in doc.tables:
        # Get the first row of the table
        first_row = table.rows[0]

        # Check if any cell in the first row is not bold
        if any(not cell.paragraphs[0].runs[0].bold for cell in first_row.cells):
            # Highlight the first row with the specified color
            for cell in first_row.cells:
                shading = parse_xml('<w:shd {} w:fill="ADD8E6"/>'.format(nsdecls('w')))
                cell._element.get_or_add_tcPr().append(shading)

        # Process each row in the table
        for row in table.rows:
            # Process each cell in the row
            for cell in row.cells:
                # Check if the cell is empty
                if cell.text.strip() == "":
                    # Set the background color of the empty cell to yellow
                    shading = parse_xml('<w:shd {} w:fill="FFFF00"/>'.format(nsdecls('w')))
                    cell._tc.get_or_add_tcPr().append(shading)


    # Save the modified document
    modified_file_path = os.path.splitext(file_path)[0] + "_modified.docx"
    doc.save(modified_file_path)

    status_label.config(text="Document modified and saved successfully.", fg="green")

    # Close the application window after completing execution
    root.destroy()

# Create the main application window
root = tk.Tk()
root.title("DOCX File Modifier")

# File Upload Section
file_label = tk.Label(root, text="Upload a DOCX file:")
file_label.pack(pady=10)

file_entry = tk.Entry(root, width=50)
file_entry.pack()

upload_button = tk.Button(root, text="Upload File", command=upload_file)
upload_button.pack(pady=5)

# Modify Button Section
modify_button = tk.Button(root, text="Modify DOCX", command=modify_docx)
modify_button.pack(pady=10)

# Status Section
status_label = tk.Label(root, text="", fg="black")
status_label.pack(pady=5)

# Start the main event loop
root.mainloop()
