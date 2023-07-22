#!/usr/bin/env python
# coding: utf-8

# In[1]:


from docx import Document


# In[2]:


pip install python-docx


# In[7]:


from docx.enum.text import WD_COLOR_INDEX
from docx.shared import RGBColor

#importing Document from docx library
ffont_name = 'Arial'
ffont_size = 12

diff_ftext = []
diff_size = []
from docx import Document

#Defining highlight color 


# Function which reads only text present in paragraph.
def read_text_from_docx(file_path):
    doc = Document(file_path)
    text = []
    for paragraph in doc.paragraphs:
        for run in paragraph.runs:
            print(run.text)
            
            font_style = run.font.bold
            font_name = run.font.name
            font_size = run.font.size
            
#             //Checking the font name if it dosn't match adding it to diff_ftext array
            if font_name != ffont_name:
                print(f"Font Name: {font_name}",run.text)
#                // try to highlight it
                run.font.highlight_color = WD_COLOR_INDEX.YELLOW
                
#                 diff_ftext.append(run.text)
                
                

#             //Checking for the font size of the element
            if font_size is not None and font_size.pt == ffont_size:
                print(f"Font Size: {font_size} and Font size is matching")
            else:
                print(f"Font size: {run.text} is not matching")
                run.font.highlight_color = WD_COLOR_INDEX.YELLOW
                diff_ftext.append(run.text)
                
                
#             //Checking the font style(B,i or U)
            if run.font.italic:
                run.font.highlight_color = WD_COLOR_INDEX.YELLOW
            if run.font.bold:
                run.font.highlight_color = WD_COLOR_INDEX.YELLOW
            if run.font.underline:
                run.font.highlight_color = WD_COLOR_INDEX.YELLOW
            else:
                print("Normal text")
                
        text.append(paragraph.text)
        
        comments_part = doc.part.comments_part

        # Work with the comments
        if comments_part:
            # Iterate over the comments
            for comment in comments_part.comments:
                print(comment.author)
                print(comment.text)
            else:
                print("No comments found in the document.")
        
#         //Removing empty string from the array
        for ele in diff_ftext:
            if ele == " ":
                diff_ftext.remove(ele)
                
    print(diff_size)
    doc.save("newoutput.docx")
    return "\n".join(text)
    

                         

# Heighlight functon
def highlight_text():
    
    doc = Document(file_path)
    for p in doc.paragraphs:
        for run in p.runs:
            for target_text in diff_ftext:
                if target_text in run.text:
                        run.font.highlight_color = WD_COLOR_INDEX.YELLOW
                        
    doc.save("newoutput.docx")
    
    
#Input
file_path = "D:\python\Demo1.docx"
document_text = read_text_from_docx(file_path)
# document_table = read_table_from_docx(file_path)

print(document_text)
# print(document_table)

print(diff_ftext)
# highlight_text()

# add_comment("Test comment")
print(diff_ftext)    


# In[5]:


doc = Document(file_path)
print(dir(doc))


# In[ ]:





# In[ ]:




