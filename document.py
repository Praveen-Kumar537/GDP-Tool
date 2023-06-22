from docx import Document

def get_run_font_properties(run):
    if run.text.strip():  # Check if run has non-empty text
        return run.font.name, run.font.size, run.font.bold
    else:
        return None, None, None

def read_table_from_docx(file_path):
    doc = Document(file_path)
    
    # Assuming the table is in the first table of the document
    table = doc.tables[0]
    
    data = []
    font_names = []
    font_sizes = []
    bold_formats = []
    
    for i, row in enumerate(table.rows):
        row_data = []
        for j, cell in enumerate(row.cells):
            if i == 0:  # First row
                cell_font_names = []
                cell_font_sizes = []
                cell_bold_formats = []
                for run in cell.paragraphs[0].runs:
                    font_name, font_size, bold_format = get_run_font_properties(run)
                    cell_font_names.append(font_name)
                    cell_font_sizes.append(font_size)
                    cell_bold_formats.append(bold_format)
                
                font_names.append(cell_font_names)
                font_sizes.append(cell_font_sizes)
                bold_formats.append(cell_bold_formats)
                
            row_data.append(cell.text)
        data.append(row_data)
    
    return data, font_names, font_sizes, bold_formats

# Provide the path to your Word document
file_path = 'sample.docx'

# Call the function to read the table and font information
table_data, font_names, font_sizes, bold_formats = read_table_from_docx(file_path)

# Print the table data
for row in table_data:
    print(row)

# Print the font information
for i in range(len(font_names[0])):
    print("Font Names (Column {}):".format(i+1), [font_name[i] for font_name in font_names])
    print("Font Sizes (Column {}):".format(i+1), [font_size[i] for font_size in font_sizes])
    print("Bold Formats (Column {}):".format(i+1), [bold_format[i] for bold_format in bold_formats])
