from docx import Document
from docx.shared import RGBColor
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml

# Load the Word document
doc = Document('sample.docx')

# Define the color for highlighting the first row (e.g., light blue)
highlight_color = RGBColor(173, 216, 230)  # Light blue highlight

# Iterate over the tables in the document
for table in doc.tables:
    # Get the first row of the table
    first_row = table.rows[0]

    # Check if any cell in the first row is not bold
    if any(not cell.paragraphs[0].runs[0].bold for cell in first_row.cells):
        # Highlight the first row with the specified color
        for cell in first_row.cells:
            shading = parse_xml('<w:shd {} w:fill="ADD8E6"/>'.format(nsdecls('w')))
            cell._element.get_or_add_tcPr().append(shading)

# Save the modified document with highlighted first rows
doc.save('modified_document.docx')
